from __future__ import annotations

import logging
from pathlib import Path

import docx
import pytest
from langchain_core.documents import Document

import src.pipeline as pipeline_module
from src.config import Settings
from src.pipeline import RAGPipeline


class FakeLLM:
    def __init__(self, model: str):
        self.model = model
        self.responses: list[object] = []
        self.prompts: list[str] = []

    def invoke(self, prompt: object) -> str:
        self.prompts.append(str(prompt))
        if self.responses:
            next_response = self.responses.pop(0)
            if isinstance(next_response, Exception):
                raise next_response
            return str(next_response)
        return "mock-response"


class FakeUploadedFile:
    def __init__(self, name: str, payload: bytes = b"dummy"):
        self.name = name
        self._payload = payload

    def getbuffer(self) -> bytes:
        return self._payload


def make_settings(tmp_path: Path, **overrides: object) -> Settings:
    base = {
        "ollama_model": "test-llm",
        "embedding_model": "test-embed",
        "rerank_model": "test-rerank",
        "chunk_size": 120,
        "chunk_overlap": 20,
        "retriever_k": 3,
        "retriever_search_type": "similarity",
        "hybrid_alpha": 0.6,
        "bm25_fetch_k": 5,
        "rerank_top_n": 2,
        "conversational_memory_turns": 4,
        "data_dir": tmp_path / "data",
        "faiss_dir": tmp_path / "faiss",
    }
    base.update(overrides)
    return Settings(**base)


def make_pipeline(monkeypatch: pytest.MonkeyPatch, tmp_path: Path, **settings_overrides: object) -> RAGPipeline:
    monkeypatch.setattr(pipeline_module, "OllamaLLM", FakeLLM)
    settings = make_settings(tmp_path, **settings_overrides)
    logger = logging.getLogger("test.pipeline")
    logger.setLevel(logging.CRITICAL)
    return RAGPipeline(settings=settings, logger=logger)


def test_load_docx_extracts_paragraphs_and_tables(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)

    sample_docx = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_paragraph("Executive summary")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "AC-1"
    table.cell(1, 1).text = "Implemented"
    doc.save(sample_docx)

    loaded = pipeline._load_docx(sample_docx)

    assert len(loaded) == 3
    assert loaded[0].page_content == "Executive summary"
    assert loaded[1].page_content == "Control | Status"
    assert loaded[2].page_content == "AC-1 | Implemented"
    assert [doc.metadata["page"] for doc in loaded] == [1, 2, 3]


def test_ingest_files_populates_required_metadata(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)

    def fake_loader(_uploaded: object) -> tuple[Path, list[Document]]:
        text = " ".join(["alpha", "beta", "gamma", "delta", "epsilon", "zeta"] * 15)
        return tmp_path / "fake.docx", [Document(page_content=text, metadata={"page": 1})]

    monkeypatch.setattr(pipeline, "_load_uploaded_documents", fake_loader)
    monkeypatch.setattr(pipeline, "_rebuild_indices", lambda *_args, **_kwargs: None)

    result = pipeline.ingest_files([FakeUploadedFile(name="audit.docx")], chunk_size=80, chunk_overlap=10)

    assert result.files_indexed == 1
    assert pipeline.document_count > 0
    first = pipeline._all_documents[0]
    assert first.metadata["source_name"] == "audit.docx"
    assert first.metadata["doc_type"] == "docx"
    assert first.metadata["chunk_id"] >= 1
    assert "uploaded_at" in first.metadata
    assert "start_index" in first.metadata


def test_ingest_files_uses_explicit_none_check_for_chunk_args(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path, chunk_size=120, chunk_overlap=20)

    captured: dict[str, int] = {}

    class FakeSplitter:
        def __init__(self, *, chunk_size: int, chunk_overlap: int, add_start_index: bool):
            captured["chunk_size"] = chunk_size
            captured["chunk_overlap"] = chunk_overlap
            captured["add_start_index"] = int(add_start_index)

        def split_documents(self, docs: list[Document]) -> list[Document]:
            return docs

    def fake_loader(_uploaded: object) -> tuple[Path, list[Document]]:
        return tmp_path / "fake.docx", [Document(page_content="abc", metadata={"page": 1})]

    monkeypatch.setattr(pipeline_module, "RecursiveCharacterTextSplitter", FakeSplitter)
    monkeypatch.setattr(pipeline, "_load_uploaded_documents", fake_loader)
    monkeypatch.setattr(pipeline, "_rebuild_indices", lambda *_args, **_kwargs: None)

    pipeline.ingest_files([FakeUploadedFile(name="a.docx")], chunk_size=0, chunk_overlap=0)
    assert captured["chunk_size"] == 0
    assert captured["chunk_overlap"] == 0

    pipeline.ingest_files([FakeUploadedFile(name="a.docx")], chunk_size=None, chunk_overlap=None)
    assert captured["chunk_size"] == pipeline.settings.chunk_size
    assert captured["chunk_overlap"] == pipeline.settings.chunk_overlap


def test_benchmark_chunk_strategies_includes_accuracy_metric(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)

    long_text = " ".join(
        [
            "incident",
            "response",
            "containment",
            "eradication",
            "recovery",
            "lessons",
            "controls",
            "risk",
            "mitigation",
            "evidence",
        ]
        * 80
    )

    def fake_loader(_uploaded: object) -> tuple[Path, list[Document]]:
        return tmp_path / "fake.pdf", [Document(page_content=long_text, metadata={"page": 1})]

    monkeypatch.setattr(pipeline, "_load_uploaded_documents", fake_loader)

    results = pipeline.benchmark_chunk_strategies(
        uploaded_file=FakeUploadedFile(name="nist.pdf"),
        chunk_sizes=[100],
        chunk_overlaps=[20],
    )

    assert len(results) == 1
    row = results[0]
    assert 0.0 <= row.accuracy_recall_at_k <= 1.0
    assert row.benchmark_queries > 0


def test_embedder_is_cached_across_pipeline_instances(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    created: list[object] = []

    class FakeEmbeddings:
        def __init__(self, **_kwargs: object):
            created.append(self)

    monkeypatch.setattr(pipeline_module, "HuggingFaceEmbeddings", FakeEmbeddings)

    first = make_pipeline(monkeypatch, tmp_path, embedding_model="cached-model")
    second = make_pipeline(monkeypatch, tmp_path, embedding_model="cached-model")

    first_embedder = first._ensure_embedder()
    second_embedder = second._ensure_embedder()

    assert len(created) == 1
    assert first_embedder is second_embedder


def test_rebuild_indices_uses_incremental_add_when_available(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)
    pipeline._all_documents = [Document(page_content="old content", metadata={"source_name": "old.docx"})]

    class FakeStore:
        def __init__(self) -> None:
            self.added_docs: list[Document] = []
            self.saved_paths: list[str] = []

        def add_documents(self, docs: list[Document]) -> None:
            self.added_docs.extend(docs)

        def save_local(self, path: str) -> None:
            self.saved_paths.append(path)

        def as_retriever(self, **kwargs: object) -> dict[str, object]:
            return {"retriever": True, "kwargs": kwargs}

    fake_store = FakeStore()
    from_documents_calls: list[int] = []

    class FakeFAISS:
        @staticmethod
        def from_documents(docs: list[Document], _embedder: object) -> FakeStore:
            from_documents_calls.append(len(docs))
            return FakeStore()

    new_doc = Document(page_content="new content", metadata={"source_name": "new.docx"})
    pipeline._all_documents.append(new_doc)
    pipeline._vector_store = fake_store

    monkeypatch.setattr(pipeline, "_ensure_embedder", lambda: object())
    monkeypatch.setattr(pipeline_module, "FAISS", FakeFAISS)

    pipeline._rebuild_indices(new_docs=[new_doc])

    assert from_documents_calls == []
    assert fake_store.added_docs == [new_doc]
    assert pipeline._retriever["retriever"] is True
    assert pipeline._bm25_index is not None


def test_rebuild_indices_falls_back_to_full_build_when_store_missing(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)
    pipeline._all_documents = [Document(page_content="doc one", metadata={})]

    from_documents_calls: list[int] = []

    class FakeStore:
        def save_local(self, _path: str) -> None:
            return

        def as_retriever(self, **_kwargs: object) -> dict[str, bool]:
            return {"ok": True}

    class FakeFAISS:
        @staticmethod
        def from_documents(docs: list[Document], _embedder: object) -> FakeStore:
            from_documents_calls.append(len(docs))
            return FakeStore()

    monkeypatch.setattr(pipeline, "_ensure_embedder", lambda: object())
    monkeypatch.setattr(pipeline_module, "FAISS", FakeFAISS)

    pipeline._rebuild_indices(new_docs=[Document(page_content="new", metadata={})])

    assert from_documents_calls == [1]
    assert pipeline._retriever == {"ok": True}


def test_rerank_returns_only_reranked_candidates(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path, rerank_top_n=2)

    class FakeCrossEncoder:
        @staticmethod
        def predict(_pairs: list[tuple[str, str]]) -> list[float]:
            return [0.2, 0.9]

    monkeypatch.setattr(pipeline, "_ensure_cross_encoder", lambda: FakeCrossEncoder())

    docs = [
        Document(page_content="doc-1", metadata={}),
        Document(page_content="doc-2", metadata={}),
        Document(page_content="doc-3", metadata={}),
    ]

    reranked = pipeline._rerank("risk controls", docs)

    assert len(reranked) == 2
    assert reranked[0].page_content == "doc-2"
    assert reranked[1].page_content == "doc-1"
    assert reranked[0].metadata["rerank_score"] == 0.9


def test_rerank_with_non_positive_topn_still_reranks_one(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path, rerank_top_n=0)

    class FakeCrossEncoder:
        @staticmethod
        def predict(_pairs: list[tuple[str, str]]) -> list[float]:
            return [0.7]

    monkeypatch.setattr(pipeline, "_ensure_cross_encoder", lambda: FakeCrossEncoder())

    docs = [Document(page_content="only-doc", metadata={})]
    reranked = pipeline._rerank("query", docs)

    assert len(reranked) == 1
    assert reranked[0].metadata["rerank_score"] == 0.7


def test_answer_self_rag_confidence_comes_from_reflection(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)
    pipeline._all_documents = [Document(page_content="seed", metadata={})]

    retrieved_docs = [Document(page_content="retrieved context", metadata={"source_name": "doc.pdf", "chunk_id": 1})]
    monkeypatch.setattr(pipeline, "_retrieve", lambda *_args, **_kwargs: retrieved_docs)

    llm = pipeline._llm
    assert isinstance(llm, FakeLLM)
    llm.responses = [
        "Base answer",
        "CONFIDENCE: 0.87\nNEEDS_REWRITE: no\nREWRITE_QUERY:\nRATIONALE: reflection ok",
    ]

    result = pipeline.answer(
        question="What is the control objective?",
        conversational=False,
        enable_self_rag=True,
    )

    assert result.answer == "Base answer"
    assert result.confidence == 0.87
    assert result.rationale == "reflection ok"


def test_answer_skips_query_rewrite_without_chat_history(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)
    pipeline._all_documents = [Document(page_content="seed", metadata={})]

    docs = [Document(page_content="retrieved context", metadata={"source_name": "doc.pdf", "chunk_id": 1})]
    monkeypatch.setattr(pipeline, "_retrieve", lambda *_args, **_kwargs: docs)

    llm = pipeline._llm
    assert isinstance(llm, FakeLLM)
    llm.responses = ["Base answer"]

    result = pipeline.answer(
        question="Follow-up question",
        conversational=True,
        chat_history=[],
    )

    assert result.used_query == "Follow-up question"
    assert len(llm.prompts) == 1
    assert "Rewrite the current user question" not in llm.prompts[0]


def test_answer_uses_query_rewrite_with_chat_history(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)
    pipeline._all_documents = [Document(page_content="seed", metadata={})]

    docs = [Document(page_content="retrieved context", metadata={"source_name": "doc.pdf", "chunk_id": 1})]
    monkeypatch.setattr(pipeline, "_retrieve", lambda *_args, **_kwargs: docs)

    llm = pipeline._llm
    assert isinstance(llm, FakeLLM)
    llm.responses = ["rewritten query", "Base answer"]

    result = pipeline.answer(
        question="And what next?",
        conversational=True,
        chat_history=[{"question": "Before", "answer": "Earlier response"}],
    )

    assert result.used_query == "rewritten query"
    assert any("Rewrite the current user question" in prompt for prompt in llm.prompts)


def test_answer_self_rag_skips_reflection_when_context_is_blank(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)
    pipeline._all_documents = [Document(page_content="seed", metadata={})]

    docs = [Document(page_content="   ", metadata={"source_name": "doc.pdf", "chunk_id": 1})]
    monkeypatch.setattr(pipeline, "_retrieve", lambda *_args, **_kwargs: docs)

    llm = pipeline._llm
    assert isinstance(llm, FakeLLM)
    llm.responses = ["Base answer"]

    result = pipeline.answer(
        question="What is this?",
        conversational=False,
        enable_self_rag=True,
    )

    assert result.answer == "Base answer"
    assert result.confidence == 0.0
    assert result.rationale == "No context to reflect on."
    assert len(llm.prompts) == 1


def test_answer_self_rag_reflection_failure_has_clear_fallback(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    pipeline = make_pipeline(monkeypatch, tmp_path)
    pipeline._all_documents = [Document(page_content="seed", metadata={})]

    retrieved_docs = [Document(page_content="retrieved context", metadata={"source_name": "doc.pdf", "chunk_id": 1})]
    monkeypatch.setattr(pipeline, "_retrieve", lambda *_args, **_kwargs: retrieved_docs)

    llm = pipeline._llm
    assert isinstance(llm, FakeLLM)
    llm.responses = ["Base answer", RuntimeError("reflection offline")]

    result = pipeline.answer(
        question="What is the control objective?",
        conversational=False,
        enable_self_rag=True,
    )

    assert result.answer == "Base answer"
    assert result.confidence == 0.6
    assert result.rationale == "Base RAG answer generated."

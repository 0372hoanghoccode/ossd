from __future__ import annotations

import os
import pickle
import re
import time
from dataclasses import dataclass
import shutil
from datetime import datetime
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any, Callable

import docx
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_ollama import OllamaLLM
from langchain_text_splitters import RecursiveCharacterTextSplitter
from rank_bm25 import BM25Okapi

from src.config import Settings
from src.corag_engine import CoRAGResult, CoRAGStep, _parse_json_robust, run_corag
from src.prompts import (
    build_corag_decompose_prompt,
    build_conversational_prompt,
    build_multi_doc_conversational_prompt,
    build_multi_doc_prompt,
    build_prompt,
    build_query_rewrite_prompt,
    build_self_reflection_prompt,
    format_chat_history,
    is_vietnamese,
    parse_reflection,
)


os.environ.setdefault("TRANSFORMERS_NO_TF", "1")
os.environ.setdefault("USE_TF", "0")
os.environ.setdefault("USE_TORCH", "1")


_EMBEDDER_CACHE: dict[str, Any] = {}


@dataclass
class IngestionResult:
    source_paths: list[Path]
    chunks: int
    seconds: float
    chunk_size: int
    chunk_overlap: int
    files_indexed: int


@dataclass
class ChunkExperimentResult:
    chunk_size: int
    chunk_overlap: int
    chunks: int
    split_seconds: float
    accuracy_recall_at_k: float
    benchmark_queries: int


@dataclass
class AnswerResult:
    answer: str
    sources: list[Document]
    retrieval_seconds: float
    generation_seconds: float
    mode: str
    confidence: float
    rationale: str
    used_query: str
    llm_total_seconds: float = 0.0
    total_seconds: float = 0.0


class RAGPipeline:
    def __init__(self, settings: Settings, logger: Any):
        self.settings = settings
        self.logger = logger
        self.settings.data_dir.mkdir(parents=True, exist_ok=True)
        self.settings.faiss_dir.mkdir(parents=True, exist_ok=True)

        self._embedder = None
        self._llm = OllamaLLM(model=self.settings.ollama_model)
        self._vector_store: FAISS | None = None
        self._retriever = None
        self._all_documents: list[Document] = []
        self._raw_documents: list[Document] = []
        self._bm25_index: BM25Okapi | None = None
        self._bm25_corpus: list[list[str]] = []
        self._cross_encoder: Any | None = None
        self._last_corag_result: CoRAGResult | None = None

        # Auto-load persisted index from disk
        self._try_load_persisted_index()

    @property
    def is_ready(self) -> bool:
        return self._retriever is not None

    @property
    def document_count(self) -> int:
        return len(self._all_documents)

    def _load_pdf(self, file_path: Path) -> list[Document]:
        loader = PDFPlumberLoader(str(file_path))
        docs = loader.load()
        # PDFPlumber uses 0-indexed pages; convert to 1-indexed for display
        for doc in docs:
            raw_page = doc.metadata.get("page", 0)
            if isinstance(raw_page, int) and raw_page >= 0:
                doc.metadata["page"] = raw_page + 1
        return docs

    def _iter_docx_blocks(self, parsed: DocxDocument) -> list[tuple[str, str]]:
        blocks: list[tuple[str, str]] = []
        for child in parsed.element.body.iterchildren():
            if isinstance(child, CT_P):
                paragraph = Paragraph(child, parsed)
                text = paragraph.text.strip()
                if text:
                    blocks.append(("paragraph", text))
                continue

            if isinstance(child, CT_Tbl):
                table = Table(child, parsed)
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if cells:
                        blocks.append(("table_row", " | ".join(cells)))

        return blocks

    def _load_docx(self, file_path: Path) -> list[Document]:
        parsed = docx.Document(str(file_path))
        blocks = self._iter_docx_blocks(parsed)

        documents: list[Document] = []
        for idx, (block_type, block_text) in enumerate(blocks, start=1):
            documents.append(
                Document(
                    page_content=block_text,
                    metadata={
                        "source": str(file_path),
                        "page": idx,
                        "docx_block_type": block_type,
                    },
                )
            )
        return documents

    @staticmethod
    def _tokenize(text: str) -> list[str]:
        return re.findall(r"\w+", text.lower())

    def _build_benchmark_queries(
        self,
        split_docs: list[Document],
        max_queries: int = 6,
    ) -> list[tuple[str, int]]:
        if not split_docs:
            return []

        if len(split_docs) <= max_queries:
            sample_indices = list(range(len(split_docs)))
        else:
            step = (len(split_docs) - 1) / max(max_queries - 1, 1)
            sample_indices = sorted({int(round(i * step)) for i in range(max_queries)})

        queries: list[tuple[str, int]] = []
        for idx in sample_indices:
            tokens = self._tokenize(split_docs[idx].page_content)
            if len(tokens) < 6:
                continue
            start = max((len(tokens) // 3) - 4, 0)
            query_tokens = tokens[start : start + 8]
            if not query_tokens:
                continue
            queries.append((" ".join(query_tokens), idx))
        return queries

    def _estimate_chunk_recall_at_k(
        self,
        split_docs: list[Document],
        top_k: int,
    ) -> tuple[float, int]:
        if not split_docs:
            return 0.0, 0

        corpus = [self._tokenize(doc.page_content) for doc in split_docs]
        if not corpus:
            return 0.0, 0

        bm25 = BM25Okapi(corpus)
        benchmark_queries = self._build_benchmark_queries(split_docs)
        if not benchmark_queries:
            return 0.0, 0

        effective_k = max(1, top_k)
        hits = 0
        evaluated = 0

        for query, target_idx in benchmark_queries:
            query_tokens = self._tokenize(query)
            if not query_tokens:
                continue
            scores = bm25.get_scores(query_tokens)
            ranked_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)
            top_indices = ranked_indices[:effective_k]
            evaluated += 1
            if target_idx in top_indices:
                hits += 1

        if evaluated == 0:
            return 0.0, 0
        return hits / evaluated, evaluated

    def _ensure_cross_encoder(self) -> Any | None:
        if self._cross_encoder is not None:
            return self._cross_encoder
        try:
            from sentence_transformers import CrossEncoder

            self._cross_encoder = CrossEncoder(self.settings.rerank_model)
        except Exception as exc:
            self.logger.warning("Cross-encoder unavailable, fallback to retrieval ranking: %s", exc)
            self._cross_encoder = None
        return self._cross_encoder

    def _ensure_embedder(self) -> Any:
        if self._embedder is not None:
            return self._embedder
        try:
            model_key = self.settings.embedding_model
            if model_key not in _EMBEDDER_CACHE:
                _EMBEDDER_CACHE[model_key] = HuggingFaceEmbeddings(
                    model_name=self.settings.embedding_model,
                    model_kwargs={"device": "cpu"},
                    encode_kwargs={"normalize_embeddings": True},
                )
            self._embedder = _EMBEDDER_CACHE[model_key]
            return self._embedder
        except ModuleNotFoundError as exc:
            if "sentence_transformers" in str(exc):
                raise ImportError(
                    "Embedding dependency is missing. Install with: pip install sentence-transformers"
                ) from exc
            raise
        except Exception as exc:
            raise RuntimeError(f"Failed to initialize embeddings: {exc}") from exc

    def _rebuild_indices(self, new_docs: list[Document] | None = None) -> None:
        if not self._all_documents:
            self._vector_store = None
            self._retriever = None
            self._bm25_index = None
            self._bm25_corpus = []
            return

        embedder = self._ensure_embedder()
        if self._vector_store is not None and new_docs:
            self._vector_store.add_documents(new_docs)
        else:
            self._vector_store = FAISS.from_documents(self._all_documents, embedder)

        self._vector_store.save_local(str(self.settings.faiss_dir))

        # Persist documents for reload after F5
        docs_path = self.settings.faiss_dir / "documents.pkl"
        with open(docs_path, "wb") as f:
            pickle.dump(self._all_documents, f)

        raw_docs_path = self.settings.faiss_dir / "raw_documents.pkl"
        with open(raw_docs_path, "wb") as f:
            pickle.dump(self._raw_documents, f)

        self._retriever = self._vector_store.as_retriever(
            search_type=self.settings.retriever_search_type,
            search_kwargs={"k": self.settings.retriever_k},
        )

        self._bm25_corpus = [self._tokenize(doc.page_content) for doc in self._all_documents]
        self._bm25_index = BM25Okapi(self._bm25_corpus) if self._bm25_corpus else None

    def _apply_metadata_filters(
        self,
        docs: list[Document],
        metadata_filters: dict[str, list[str]] | None,
    ) -> list[Document]:
        if not metadata_filters:
            return docs

        source_filter = set(metadata_filters.get("source_name", []))
        type_filter = set(metadata_filters.get("doc_type", []))
        upload_filter = set(metadata_filters.get("uploaded_at", []))

        filtered: list[Document] = []
        for doc in docs:
            source_name = str(doc.metadata.get("source_name", ""))
            doc_type = str(doc.metadata.get("doc_type", ""))
            uploaded_at = str(doc.metadata.get("uploaded_at", ""))
            source_ok = not source_filter or source_name in source_filter
            type_ok = not type_filter or doc_type in type_filter
            upload_ok = not upload_filter or uploaded_at in upload_filter
            if source_ok and type_ok and upload_ok:
                filtered.append(doc)
        return filtered

    def _vector_retrieve(
        self,
        query: str,
        fetch_k: int,
        metadata_filters: dict[str, list[str]] | None,
    ) -> list[Document]:
        if not self._vector_store:
            return []

        pairs = self._vector_store.similarity_search_with_relevance_scores(query, k=max(fetch_k, 1))
        docs: list[Document] = []
        for doc, score in pairs:
            doc.metadata["retrieval_score"] = float(score)
            docs.append(doc)

        filtered = self._apply_metadata_filters(docs, metadata_filters)
        return filtered[: self.settings.retriever_k]

    def _keyword_retrieve(
        self,
        query: str,
        fetch_k: int,
        metadata_filters: dict[str, list[str]] | None,
    ) -> list[Document]:
        if not self._bm25_index:
            return []

        query_tokens = self._tokenize(query)
        if not query_tokens:
            return []

        scores = self._bm25_index.get_scores(query_tokens)
        indexed_scores = sorted(enumerate(scores), key=lambda item: item[1], reverse=True)[:fetch_k]
        docs: list[Document] = []
        for idx, score in indexed_scores:
            doc = self._all_documents[idx]
            copy_doc = Document(page_content=doc.page_content, metadata=dict(doc.metadata))
            copy_doc.metadata["retrieval_score"] = float(score)
            docs.append(copy_doc)

        filtered = self._apply_metadata_filters(docs, metadata_filters)
        return filtered[: self.settings.retriever_k]

    def _hybrid_retrieve(
        self,
        query: str,
        fetch_k: int,
        metadata_filters: dict[str, list[str]] | None,
    ) -> list[Document]:
        vector_docs = self._vector_retrieve(query, fetch_k, metadata_filters)
        keyword_docs = self._keyword_retrieve(query, fetch_k, metadata_filters)

        merged: dict[str, tuple[Document, float]] = {}

        def doc_key(doc: Document) -> str:
            return (
                f"{doc.metadata.get('source_name', '')}|"
                f"{doc.metadata.get('chunk_id', '')}|"
                f"{doc.metadata.get('start_index', '')}"
            )

        for rank, doc in enumerate(vector_docs):
            score = self.settings.hybrid_alpha * (1.0 / (rank + 1))
            merged[doc_key(doc)] = (doc, score)

        for rank, doc in enumerate(keyword_docs):
            key = doc_key(doc)
            score = (1.0 - self.settings.hybrid_alpha) * (1.0 / (rank + 1))
            if key in merged:
                current_doc, current_score = merged[key]
                merged[key] = (current_doc, current_score + score)
            else:
                merged[key] = (doc, score)

        reranked = sorted(merged.values(), key=lambda item: item[1], reverse=True)[: self.settings.retriever_k]
        docs = [doc for doc, score in reranked]
        for doc, score in reranked:
            doc.metadata["retrieval_score"] = float(score)
        return docs

    def _rerank(self, query: str, docs: list[Document]) -> list[Document]:
        if not docs:
            return docs
        cross_encoder = self._ensure_cross_encoder()
        if cross_encoder is None:
            return docs

        top_n = min(max(self.settings.rerank_top_n, 1), len(docs))
        candidates = docs[:top_n]
        pairs = [(query, doc.page_content) for doc in candidates]
        scores = cross_encoder.predict(pairs)
        ranked = sorted(zip(candidates, scores), key=lambda item: item[1], reverse=True)

        reranked_docs = [doc for doc, _ in ranked]
        for doc, score in ranked:
            doc.metadata["rerank_score"] = float(score)
        return reranked_docs

    def decompose_question(self, question: str) -> list[str]:
        """
        Decompose question into atomic retrievable parts for CoRAG.
        Never raises and always returns at least one part.
        """
        safe_question = question.strip() or question

        def dedupe_parts(parts: list[str]) -> list[str]:
            cleaned: list[str] = []
            seen: set[str] = set()
            for part in parts:
                text = str(part).strip()
                if not text:
                    continue
                key = text.lower()
                if key in seen:
                    continue
                seen.add(key)
                cleaned.append(text)
            return cleaned[:5]

        try:
            prompt = build_corag_decompose_prompt(safe_question)
            raw = str(self._llm.invoke(prompt)).strip()
            parsed = _parse_json_robust(raw)
            raw_parts = parsed.get("parts", [])
            if isinstance(raw_parts, list):
                parts = dedupe_parts([str(item) for item in raw_parts])
                if parts:
                    return parts
        except Exception as exc:
            self.logger.warning("CoRAG question decomposition failed, fallback to heuristic split: %s", exc)

        try:
            split_candidates = re.split(r"\s*(?:,|;|\band\b|\bvà\b|\?|\.)\s*", safe_question, flags=re.IGNORECASE)
            heuristic = dedupe_parts(split_candidates)
            if heuristic:
                return heuristic
        except Exception:
            pass

        return [safe_question] if safe_question else [question]

    def _detect_file_type(self, uploaded_file: Any) -> tuple[str, str]:
        filename = (uploaded_file.name or "").lower()
        if filename.endswith(".pdf"):
            return ".pdf", "pdf"
        if filename.endswith(".docx"):
            return ".docx", "docx"
        raise ValueError("Only .pdf and .docx files are supported.")

    def _load_uploaded_documents(self, uploaded_file: Any) -> tuple[Path, list[Document]]:
        suffix, file_type = self._detect_file_type(uploaded_file)

        with NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(uploaded_file.getbuffer())
            temp_path = Path(temp_file.name)

        if file_type == "pdf":
            docs = self._load_pdf(temp_path)
        else:
            docs = self._load_docx(temp_path)

        return temp_path, docs

    def ingest_file(
        self,
        uploaded_file: Any,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
    ) -> IngestionResult:
        return self.ingest_files([uploaded_file], chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    def ingest_files(
        self,
        uploaded_files: list[Any],
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
    ) -> IngestionResult:
        started = time.perf_counter()

        effective_chunk_size = chunk_size if chunk_size is not None else self.settings.chunk_size
        effective_chunk_overlap = chunk_overlap if chunk_overlap is not None else self.settings.chunk_overlap

        all_new_docs: list[Document] = []
        source_paths: list[Path] = []
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=effective_chunk_size,
            chunk_overlap=effective_chunk_overlap,
            add_start_index=True,
        )

        for uploaded_file in uploaded_files:
            temp_path, docs = self._load_uploaded_documents(uploaded_file)
            source_paths.append(temp_path)

            source_name = uploaded_file.name or temp_path.name
            uploaded_at = datetime.now().isoformat(timespec="seconds")
            doc_type = source_name.split(".")[-1].lower() if "." in source_name else "unknown"

            for d in docs:
                d.metadata["source_name"] = source_name
                d.metadata["uploaded_at"] = uploaded_at
                d.metadata["doc_type"] = doc_type
                self._raw_documents.append(d)

            split_docs = splitter.split_documents(docs)

            for idx, doc in enumerate(split_docs, start=1):
                doc.metadata["chunk_id"] = idx
                doc.metadata["source_name"] = source_name
                doc.metadata["uploaded_at"] = uploaded_at
                doc.metadata["doc_type"] = doc_type
                doc.metadata.setdefault("start_index", -1)
                all_new_docs.append(doc)

        self._all_documents.extend(all_new_docs)
        self._rebuild_indices(new_docs=all_new_docs)

        elapsed = time.perf_counter() - started
        self.logger.info(
            "Indexed %s chunks from %s file(s) in %.2fs",
            len(all_new_docs),
            len(uploaded_files),
            elapsed,
        )

        return IngestionResult(
            source_paths=source_paths,
            chunks=len(all_new_docs),
            seconds=elapsed,
            chunk_size=effective_chunk_size,
            chunk_overlap=effective_chunk_overlap,
            files_indexed=len(uploaded_files),
        )

    def benchmark_single_config(
        self,
        uploaded_files: list[Any],
        question: str,
        chunk_size: int,
        chunk_overlap: int,
        retrieval_mode: str = "vector",
    ) -> dict:
        """Run a full RAG answer for a single chunk config over multiple files."""
        import time as _time

        # Load and split
        all_docs = []
        source_names = []
        for file in uploaded_files:
            temp_path, docs = self._load_uploaded_documents(file)
            source_name = file.name or temp_path.name
            source_names.append(source_name)
            for d in docs:
                d.metadata["source_name"] = source_name
            all_docs.extend(docs)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            add_start_index=True,
        )
        split_docs = splitter.split_documents(all_docs)

        # Build a temporary FAISS index
        embedder = self._ensure_embedder()
        for idx, doc in enumerate(split_docs, start=1):
            doc.metadata["chunk_id"] = idx
            doc.metadata.setdefault("start_index", -1)

        tmp_store = FAISS.from_documents(split_docs, embedder)

        # Retrieve
        retrieve_t0 = _time.perf_counter()
        pairs = tmp_store.similarity_search_with_relevance_scores(question, k=self.settings.retriever_k)
        retrieved_docs = [doc for doc, _ in pairs]
        retrieval_sec = _time.perf_counter() - retrieve_t0

        # Generate answer
        context = "\n\n".join(doc.page_content for doc in retrieved_docs)
        from src.prompts import build_prompt
        prompt = build_prompt(context, question)
        gen_t0 = _time.perf_counter()
        answer = str(self._llm.invoke(prompt)).strip()
        generation_sec = _time.perf_counter() - gen_t0

        return {
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
            "total_chunks": len(split_docs),
            "retrieved_docs": len(retrieved_docs),
            "retrieval_ms": round(retrieval_sec * 1000),
            "generation_ms": round(generation_sec * 1000),
            "total_ms": round((retrieval_sec + generation_sec) * 1000),
            "answer": answer,
            "question": question,
            "source": ", ".join(source_names),
        }

    def benchmark_chunk_strategies(
        self,
        uploaded_files: list[Any],
        chunk_sizes: list[int],
        chunk_overlaps: list[int],
    ) -> list[ChunkExperimentResult]:
        all_docs = []
        for file in uploaded_files:
            _, docs = self._load_uploaded_documents(file)
            all_docs.extend(docs)
        
        results: list[ChunkExperimentResult] = []
        for chunk_size in chunk_sizes:
            for chunk_overlap in chunk_overlaps:
                started = time.perf_counter()
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    add_start_index=True,
                )
                split_docs = splitter.split_documents(all_docs)
                elapsed = time.perf_counter() - started
                accuracy_recall_at_k, benchmark_queries = self._estimate_chunk_recall_at_k(
                    split_docs, top_k=self.settings.retriever_k,
                )
                results.append(ChunkExperimentResult(
                    chunk_size=chunk_size, chunk_overlap=chunk_overlap,
                    chunks=len(split_docs), split_seconds=elapsed,
                    accuracy_recall_at_k=accuracy_recall_at_k, benchmark_queries=benchmark_queries,
                ))
        return results

    def clear_vector_store(self) -> None:
        self._vector_store = None
        self._retriever = None
        self._all_documents = []
        self._raw_documents = []
        self._bm25_index = None
        self._bm25_corpus = []
        if self.settings.faiss_dir.exists():
            shutil.rmtree(self.settings.faiss_dir)
        self.settings.faiss_dir.mkdir(parents=True, exist_ok=True)

    def remove_source(self, source_name: str) -> int:
        """Remove all chunks belonging to source_name, rebuild indices."""
        before = len(self._all_documents)
        self._all_documents = [d for d in self._all_documents if str(d.metadata.get("source_name", "")) != source_name]
        self._raw_documents = [d for d in self._raw_documents if str(d.metadata.get("source_name", "")) != source_name]
        removed = before - len(self._all_documents)
        if removed > 0:
            self._vector_store = None
            self._retriever = None
            self._bm25_index = None
            self._bm25_corpus = []
            if self._all_documents:
                self._rebuild_indices()
            else:
                if self.settings.faiss_dir.exists():
                    shutil.rmtree(self.settings.faiss_dir)
                self.settings.faiss_dir.mkdir(parents=True, exist_ok=True)
        return removed

    def _try_load_persisted_index(self) -> None:
        """Try to load FAISS index and documents from disk (survives F5 refresh)."""
        faiss_index_file = self.settings.faiss_dir / "index.faiss"
        docs_pickle = self.settings.faiss_dir / "documents.pkl"
        raw_docs_pickle = self.settings.faiss_dir / "raw_documents.pkl"
        if not faiss_index_file.exists() or not docs_pickle.exists():
            return
        try:
            embedder = self._ensure_embedder()
            self._vector_store = FAISS.load_local(str(self.settings.faiss_dir), embedder, allow_dangerous_deserialization=True)
            with open(docs_pickle, "rb") as f:
                self._all_documents = pickle.load(f)
            if raw_docs_pickle.exists():
                with open(raw_docs_pickle, "rb") as f:
                    self._raw_documents = pickle.load(f)
            self._retriever = self._vector_store.as_retriever(search_type=self.settings.retriever_search_type, search_kwargs={"k": self.settings.retriever_k})
            self._bm25_corpus = [self._tokenize(doc.page_content) for doc in self._all_documents]
            self._bm25_index = BM25Okapi(self._bm25_corpus) if self._bm25_corpus else None
            self.logger.info("Loaded persisted index: %d documents from %s", len(self._all_documents), self.settings.faiss_dir)
        except Exception as exc:
            self.logger.warning("Failed to load persisted index, starting fresh: %s", exc)
            self._vector_store = None
            self._retriever = None
            self._all_documents = []

    def list_available_sources(self) -> dict[str, list[str]]:
        source_names = sorted({str(doc.metadata.get("source_name", "")) for doc in self._all_documents if doc.metadata.get("source_name")})
        doc_types = sorted({str(doc.metadata.get("doc_type", "")) for doc in self._all_documents if doc.metadata.get("doc_type")})
        uploaded_at = sorted({str(doc.metadata.get("uploaded_at", "")) for doc in self._all_documents if doc.metadata.get("uploaded_at")})
        return {"source_name": source_names, "doc_type": doc_types, "uploaded_at": uploaded_at}

    def list_indexed_documents(self) -> list[dict[str, Any]]:
        """Return per-file metadata summary of all indexed documents."""
        file_map: dict[str, dict[str, Any]] = {}
        for doc in self._all_documents:
            source_name = str(doc.metadata.get("source_name", "unknown"))
            if source_name not in file_map:
                file_map[source_name] = {"source_name": source_name, "doc_type": str(doc.metadata.get("doc_type", "?")), "uploaded_at": str(doc.metadata.get("uploaded_at", "?")), "chunks": 0, "pages": set()}
            file_map[source_name]["chunks"] += 1
            page = doc.metadata.get("page")
            if page is not None:
                file_map[source_name]["pages"].add(page)
        result: list[dict[str, Any]] = []
        for info in file_map.values():
            result.append({"Tên file": info["source_name"], "Loại": info["doc_type"].upper(), "Ngày upload": info["uploaded_at"], "Số chunks": info["chunks"], "Số trang": len(info["pages"])})
        return result

    def _retrieve(self, query: str, mode: str, metadata_filters: dict[str, list[str]] | None) -> list[Document]:
        fetch_k = max(self.settings.bm25_fetch_k, self.settings.retriever_k)
        if mode == "keyword":
            return self._keyword_retrieve(query, fetch_k, metadata_filters)
        if mode == "hybrid":
            return self._hybrid_retrieve(query, fetch_k, metadata_filters)
        return self._vector_retrieve(query, fetch_k, metadata_filters)

    def answer(
        self,
        question: str,
        chat_history: list[dict[str, str]] | None = None,
        retrieval_mode: str = "vector",
        metadata_filters: dict[str, list[str]] | None = None,
        enable_rerank: bool = False,
        enable_self_rag: bool = False,
        conversational: bool = True,
        enable_query_rewrite: bool = True,
    ) -> AnswerResult:
        if not self._all_documents:
            raise ValueError("Retriever is not initialized. Upload and process a document first.")

        overall_started = time.perf_counter()
        llm_total_elapsed = 0.0

        def _invoke_llm(prompt: object) -> object:
            nonlocal llm_total_elapsed
            llm_started = time.perf_counter()
            result = self._llm.invoke(prompt)
            llm_total_elapsed += time.perf_counter() - llm_started
            return result

        chat_history = chat_history or []
        history_text = format_chat_history(chat_history, self.settings.conversational_memory_turns)

        used_query = question
        if conversational and chat_history and enable_query_rewrite:
            try:
                rewrite_prompt = build_query_rewrite_prompt(question=question, chat_history=history_text)
                rewritten = str(_invoke_llm(rewrite_prompt)).strip()
                if rewritten:
                    used_query = rewritten
            except Exception as exc:
                self.logger.warning("Query rewrite failed, fallback to original query: %s", exc)

        retrieve_started = time.perf_counter()
        docs = self._retrieve(used_query, retrieval_mode, metadata_filters)
        if enable_rerank:
            docs = self._rerank(used_query, docs)
        retrieval_elapsed = time.perf_counter() - retrieve_started

        if not docs:
            fallback_answer = (
                "Không tìm thấy ngữ cảnh phù hợp với bộ lọc hiện tại."
                if is_vietnamese(question)
                else "No relevant context found for current filters."
            )
            return AnswerResult(
                answer=fallback_answer,
                sources=[],
                retrieval_seconds=retrieval_elapsed,
                generation_seconds=0.0,
                mode=retrieval_mode,
                confidence=0.0,
                rationale="No context after retrieval/filtering.",
                used_query=used_query,
                llm_total_seconds=llm_total_elapsed,
                total_seconds=time.perf_counter() - overall_started,
            )

        context = "\n\n".join(doc.page_content for doc in docs)
        if conversational:
            prompt = build_multi_doc_conversational_prompt(docs=docs, question=question, chat_history=history_text)
        else:
            prompt = build_multi_doc_prompt(docs=docs, question=question)

        generate_started = time.perf_counter()
        answer = _invoke_llm(prompt)
        generation_elapsed = time.perf_counter() - generate_started

        final_answer = str(answer).strip()
        confidence: float | None = None
        rationale = ""

        if enable_self_rag:
            if not context.strip():
                confidence = 0.0
                rationale = "No context to reflect on."
            else:
                try:
                    reflection_prompt = build_self_reflection_prompt(question=question, answer=final_answer, context=context)
                    reflection_text = str(_invoke_llm(reflection_prompt))
                    confidence, needs_rewrite, rewrite_query, rationale = parse_reflection(reflection_text)
                    rationale = rationale or "Self-RAG reflection completed."

                    if needs_rewrite and rewrite_query and rewrite_query.lower() != used_query.lower():
                        used_query = rewrite_query
                        second_docs = self._retrieve(used_query, retrieval_mode, metadata_filters)
                        if enable_rerank:
                            second_docs = self._rerank(used_query, second_docs)
                        if second_docs:
                            second_context = "\n\n".join(doc.page_content for doc in second_docs)
                            second_prompt = (
                                build_conversational_prompt(second_context, question, history_text)
                                if conversational
                                else build_prompt(second_context, question)
                            )
                            second_answer = str(_invoke_llm(second_prompt)).strip()
                            second_reflection = str(
                                _invoke_llm(
                                    build_self_reflection_prompt(
                                        question=question,
                                        answer=second_answer,
                                        context=second_context,
                                    )
                                )
                            )
                            second_confidence, _, _, second_rationale = parse_reflection(second_reflection)
                            if second_confidence >= (confidence if confidence is not None else 0.0):
                                final_answer = second_answer
                                docs = second_docs
                                confidence = second_confidence
                                rationale = second_rationale or rationale
                except Exception as exc:
                    self.logger.warning("Self-RAG reflection failed: %s", exc)
                    confidence = 0.6
                    rationale = "Base RAG answer generated."
        else:
            confidence = 0.6
            rationale = "Base RAG answer generated."

        if confidence is None:
            confidence = 0.6
        if not rationale:
            rationale = "Base RAG answer generated."

        self.logger.info(
            "Answered query in %.2fs (retrieve=%.2fs, generate=%.2fs)",
            retrieval_elapsed + generation_elapsed,
            retrieval_elapsed,
            generation_elapsed,
        )

        return AnswerResult(
            answer=final_answer,
            sources=docs,
            retrieval_seconds=retrieval_elapsed,
            generation_seconds=generation_elapsed,
            mode=retrieval_mode,
            confidence=confidence,
            rationale=rationale,
            used_query=used_query,
            llm_total_seconds=llm_total_elapsed,
            total_seconds=time.perf_counter() - overall_started,
        )

    def answer_corag(
        self,
        question: str,
        max_steps: int = 4,
        retrieval_mode: str = "vector",
        metadata_filters: dict[str, list[str]] | None = None,
        enable_rerank: bool = False,
        step_callback: Callable[[CoRAGStep], None] | None = None,
    ) -> AnswerResult:
        """
        Run CoRAG loop using existing retrieval infrastructure.
        Never raises: on errors, returns a fallback AnswerResult.
        """
        started = time.perf_counter()
        self._last_corag_result = None

        try:
            if not self._all_documents:
                no_data_answer = (
                    "Retriever chưa sẵn sàng. Vui lòng upload và index tài liệu trước."
                    if is_vietnamese(question)
                    else "Retriever is not initialized. Upload and process a document first."
                )
                return AnswerResult(
                    answer=no_data_answer,
                    sources=[],
                    retrieval_seconds=0.0,
                    generation_seconds=0.0,
                    mode=f"corag-{retrieval_mode}",
                    confidence=0.0,
                    rationale="No indexed documents available.",
                    used_query=question,
                )

            required_parts = self.decompose_question(question)

            def retrieve_fn(query: str, k: int) -> list[Document]:
                effective_fetch = max(k, self.settings.bm25_fetch_k)

                if retrieval_mode == "hybrid":
                    docs = self._hybrid_retrieve(query, effective_fetch, metadata_filters)
                elif retrieval_mode == "keyword":
                    docs = self._keyword_retrieve(query, effective_fetch, metadata_filters)
                else:
                    docs = self._vector_retrieve(query, effective_fetch, metadata_filters)

                if enable_rerank:
                    docs = self._rerank(query, docs)
                return docs[: max(k, 1)]

            def llm_invoke_fn(prompt: str) -> str:
                raw = str(self._llm.invoke(prompt)).strip()
                raw = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL).strip()
                return raw

            corag_result = run_corag(
                question=question,
                retrieve_fn=retrieve_fn,
                llm_invoke_fn=llm_invoke_fn,
                required_parts=required_parts,
                max_steps=max_steps,
                step_k=self.settings.retriever_k,
                step_callback=step_callback,
            )
            self._last_corag_result = corag_result

            seen_sources: set[str] = set()
            unique_sources: list[Document] = []
            for doc in corag_result.all_docs:
                key = (
                    f"{doc.metadata.get('source_name', '')}|"
                    f"{doc.metadata.get('chunk_id', '')}|"
                    f"{doc.metadata.get('start_index', '')}|"
                    f"{hash(doc.page_content)}"
                )
                if key in seen_sources:
                    continue
                seen_sources.add(key)
                unique_sources.append(doc)

            elapsed = time.perf_counter() - started
            return AnswerResult(
                answer=corag_result.answer,
                sources=unique_sources[: self.settings.retriever_k],
                retrieval_seconds=elapsed,
                generation_seconds=0.0,
                mode=f"corag-{retrieval_mode}",
                confidence=1.0 if corag_result.sufficient else 0.5,
                rationale=f"{corag_result.steps} steps, {corag_result.total_docs} unique docs",
                used_query=question,
            )
        except Exception as exc:
            self.logger.exception("CoRAG answering failed")
            fallback = (
                f"CoRAG gặp lỗi và không thể hoàn tất: {exc}"
                if is_vietnamese(question)
                else f"CoRAG failed to complete: {exc}"
            )
            elapsed = time.perf_counter() - started
            return AnswerResult(
                answer=fallback,
                sources=[],
                retrieval_seconds=elapsed,
                generation_seconds=0.0,
                mode=f"corag-{retrieval_mode}",
                confidence=0.0,
                rationale="CoRAG execution fallback due to runtime error.",
                used_query=question,
            )
